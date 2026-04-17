# -*- coding: utf-8 -*-
"""NPLatform 전략 보고서 — McKinsey-style Word 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── McKinsey 색상 팔레트 ───────────────────────────────
NAVY = RGBColor(0x00, 0x1E, 0x62)        # 메인 네이비
BLUE = RGBColor(0x00, 0x6B, 0xB6)        # 액센트 블루
GRAY_DARK = RGBColor(0x3C, 0x3C, 0x3C)   # 본문
GRAY_MID = RGBColor(0x80, 0x80, 0x80)    # 서브
GRAY_LIGHT = RGBColor(0xE6, 0xE6, 0xE6)  # 테이블 헤더
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xC0, 0x00, 0x00)

FONT_KR = 'Malgun Gothic'
FONT_EN = 'Calibri'

doc = Document()

# ─── 기본 스타일 설정 ─────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT_KR
style.font.size = Pt(10)
style.font.color.rgb = GRAY_DARK
rPr = style.element.rPr
rFonts = rPr.rFonts if rPr.rFonts is not None else OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), FONT_KR)
rFonts.set(qn('w:ascii'), FONT_EN)
rFonts.set(qn('w:hAnsi'), FONT_EN)
if rPr.rFonts is None:
    rPr.append(rFonts)

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# ─── 헬퍼 함수 ─────────────────────────────────────
def set_run(run, size=10, bold=False, color=GRAY_DARK, font=FONT_KR):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_KR)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rPr.append(rFonts)

def add_para(text, size=10, bold=False, color=GRAY_DARK, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p

def add_h1(text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    if number:
        run1 = p.add_run(f"{number:02d}  ")
        set_run(run1, size=16, bold=True, color=BLUE)
    run2 = p.add_run(text)
    set_run(run2, size=16, bold=True, color=NAVY)
    # 하단 라인
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '001E62')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    # 블루 세로 바 느낌 (▌ 사용)
    run1 = p.add_run("▌ ")
    set_run(run1, size=12, bold=True, color=BLUE)
    run2 = p.add_run(text)
    set_run(run2, size=12, bold=True, color=NAVY)

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=11, bold=True, color=BLUE)

def add_bullet(text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    # 기존 runs 제거
    for r in list(p.runs):
        r.text = ''
    run = p.add_run(text)
    set_run(run, size=size, color=GRAY_DARK)

def add_key_takeaway(text):
    """상단 핵심 메시지 박스 (McKinsey governing thought)"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '001E62')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("KEY TAKEAWAY   ")
    set_run(run, size=9, bold=True, color=RGBColor(0x60, 0xA5, 0xFA))
    run2 = p.add_run(text)
    set_run(run2, size=11, bold=True, color=WHITE)
    doc.add_paragraph()  # 간격

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '001E62')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=9, bold=True, color=WHITE)
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F4F6FA')
                tcPr.append(shd)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            set_run(run, size=9, color=GRAY_DARK)
    doc.add_paragraph()

def add_page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("STRATEGIC TRANSFORMATION REPORT")
set_run(run, size=11, bold=True, color=BLUE, font=FONT_EN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("NPLatform")
set_run(run, size=36, bold=True, color=NAVY, font=FONT_EN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("규제 준수형 NPL 거래 인프라로의 전면 전환")
set_run(run, size=16, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("From Information Portal to Transaction Infrastructure")
set_run(run, size=11, color=GRAY_MID, font=FONT_EN)

for _ in range(6):
    doc.add_paragraph()

# 하단 박스
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
cell.width = Cm(14)
tc = cell._tc
tcPr = tc.get_or_add_tcPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F4F6FA')
tcPr.append(shd)
for label, value in [
    ("Engagement Period", "30 Business Days (6 Phases)"),
    ("Target Outcome", "연 매출 920억 · 매각자 락인율 85%"),
    ("Regulatory Scope", "개인정보보호법 · 신용정보법 · 전자금융거래법"),
    ("Version", "v4.0 (2026-04-07)"),
]:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}   ")
    set_run(r1, size=9, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run(r2, size=9, color=GRAY_DARK)
# 빈 첫 단락 제거
cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

add_page_break()

# ════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════
add_h1("Executive Summary", number=1)

add_key_takeaway(
    "NPLatform은 '정보 포털'에서 '거래 인프라'로 전환하여, "
    "규제 준수 기반 4단계 정보 공개 모델과 자동 마스킹 파이프라인을 통해 "
    "연 920억 매출과 85% 매각자 락인율을 달성한다."
)

add_h2("Situation — 현재 상황")
add_bullet("국내 NPL 시장은 약 30조원 규모이나 거래가 금융기관 간 폐쇄 네트워크에 집중되어 있음")
add_bullet("기존 플랫폼들은 '정보 제공'에 머물러 매각자·매수자가 연결된 뒤 플랫폼을 이탈(disintermediation)")
add_bullet("채무자 개인정보·신용정보 노출 이슈로 금감원·금융위 규제 대응이 플랫폼화의 핵심 장벽")

add_h2("Complication — 전환이 시급한 이유")
add_bullet("정보 비대칭을 무기로 한 기존 모델은 공급자(AMC·금융기관)의 정보 우위 앞에서 가치 창출 실패")
add_bullet("개인정보 보호 리스크가 미해결 상태에서는 기관 고객 확보 및 투자 유치가 구조적으로 불가능")
add_bullet("수수료 상한 규제 환경(양측 각 0.9% 이내)에서 단순 중개 모델로는 수익성 확보 어려움")

add_h2("Resolution — 핵심 권고 (Governing Thought)")
add_para(
    "\"담보는 열고, 사람은 가린다. 거래는 묶고, 정보는 지킨다.\"",
    size=12, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10
)
add_bullet("핵심 투자판단 정보(채권잔액·매각희망가·감정평가서)를 L0 공개로 전환 → 트래픽·SEO 유입 극대화")
add_bullet("채무자 개인정보는 자동 마스킹 + 계약 체결 후 법적 절차로만 이전 → 규제 리스크 제로")
add_bullet("NDA·LOI·에스크로·계약서 전자화로 거래 실행 인프라 독점 → 매각자·매수자 양방향 락인")
add_bullet("자료 완성도 점수(0~10)로 금융기관 간 자료 품질 경쟁 유도 → 플랫폼 자산 자체가 경쟁력")

add_h2("Expected Impact — 12개월 목표")
add_table(
    ["KPI", "목표", "근거"],
    [
        ["월간 상장 건수", "100건+", "저축은행·AMC 전속계약 30곳"],
        ["거래 완결률", "35%+", "에스크로 의무화 + NDA 락인"],
        ["매각자 락인율 (재상장)", "85%+", "업계 최저 0.5% + AI 평가툴"],
        ["연 거래 총액", "6,000억+", "평균 딜 50억 × 월 100건"],
        ["연 매출", "920억+", "수수료·구독·에스크로·SaaS 합산"],
        ["개인정보 침해 사고", "0건", "자동 마스킹 + PII 분리 저장"],
    ],
    col_widths=[4, 3, 9]
)

add_page_break()

# ════════════════════════════════════════════════════
# 2. 현황 진단 (Diagnostic)
# ════════════════════════════════════════════════════
add_h1("Current State Diagnostic", number=2)

add_h2("2.1  개발 현황 (As-Is)")
add_table(
    ["영역", "현재 상태", "전략 v4 적합성"],
    [
        ["/exchange 목록·상세", "구현 완료", "중 — 필드 확장 및 티어 게이팅 필요"],
        ["/exchange/demands (매수 수요)", "구현 완료", "상"],
        ["/exchange/institutions", "구현 완료", "상"],
        ["티어 시스템 (L0~L3)", "미구현", "최우선 과제"],
        ["자동 마스킹 파이프라인", "미구현", "최우선 과제"],
        ["NDA·LOI 전자서명", "미구현", "최우선 과제"],
        ["전문투자자 인증(KYC)", "미구현", "우선"],
        ["PII 분리 저장 DB", "미구현", "우선"],
        ["자료 완성도 점수", "미구현", "우선"],
        ["관리자 마스킹 검수", "미구현", "우선"],
    ],
    col_widths=[5, 4, 7]
)

add_h2("2.2  4대 구조적 갭 (MECE)")

add_h3("Gap 1.  규제 준수 기반 부재")
add_bullet("개인정보·신용정보 관련 자동 처리 체계가 없어 금융기관 고객 확보 불가")
add_bullet("DPO 조직·ISMS-P 인증·개인정보 영향평가 등 필수 요건 미충족")

add_h3("Gap 2.  티어 기반 정보 공개 모델 부재")
add_bullet("현재는 모든 사용자가 동일 정보에 접근 → 트래픽 유입·락인 기제 없음")
add_bullet("게이팅(NDA/LOI)을 통한 거래 의사 스크리닝 불가")

add_h3("Gap 3.  거래 실행 인프라 부재")
add_bullet("NDA·LOI·에스크로·계약서 전자화 기능 없음 → 거래가 플랫폼 외부에서 완결됨")
add_bullet("감사 로그·비우회 조항 등 법적 보호 장치 부재")

add_h3("Gap 4.  정보 품질 차별화 장치 부재")
add_bullet("금융기관별 자료 수준 격차가 그대로 노출되어 매수자 신뢰 저하")
add_bullet("매각자에게 자료 보충 인센티브가 없어 품질 개선 동기 부재")

add_page_break()

# ════════════════════════════════════════════════════
# 3. 전략 프레임워크 (Strategic Framework)
# ════════════════════════════════════════════════════
add_h1("Strategic Framework", number=3)

add_h2("3.1  5대 전략 축")
add_table(
    ["Pillar", "핵심 원리", "실행 메커니즘"],
    [
        ["① Graduated Disclosure", "티어별 정보 공개", "L0→L1→L2→L3 단계별 해금"],
        ["② Regulatory Compliance", "개인정보 완전 보호", "자동 마스킹 + PII 분리 저장"],
        ["③ End-to-End Process", "거래 전 과정 내재화", "NDA→LOI→에스크로→계약"],
        ["④ Institution Lock-in", "공급자 전속화", "최저 수수료 + AI 평가 + B2B SaaS"],
        ["⑤ Legal Safeguards", "우회 방지 법적 장치", "비우회 조항·감사로그·위약금"],
    ],
    col_widths=[4.5, 4, 7.5]
)

add_h2("3.2  4단계 정보 공개 모델")
add_para(
    "McKinsey 7-Step 공개 전략을 NPL 거래 특성에 맞춰 4단계로 재구성한다. "
    "핵심은 '판단에 필요한 정보'는 공개하되 '거래 실행에 필요한 권한'을 유료화하는 것이다.",
    size=10, space_after=6
)

add_table(
    ["Tier", "진입 조건", "공개 필드 (누적)", "목적"],
    [
        ["L0", "무료·무로그인",
         "채권잔액, 매각희망가, 할인율, 감정평가 요약, 담보종류, 지역(읍/면/동), 채무자유형",
         "트래픽·SEO 유입"],
        ["L1", "본인인증",
         "감정평가서 PDF(마스킹), 등기요약, 권리관계, 임차요약, 유사 낙찰가, Q&A",
         "회원 풀 확보"],
        ["L2", "전문투자자 + NDA",
         "등기 원본, 실사 사진, 재무제표(법인), 채무자 통계(비식별), 매각자 기관명",
         "진지한 매수자 스크리닝"],
        ["L3", "LOI + 매각자 승인",
         "매각자 공식 채널, 실사 원본, 협상 테이블, 에스크로 개설, 계약서",
         "거래 완결 (플랫폼 락인)"],
    ],
    col_widths=[1.5, 3.5, 7.5, 3.5]
)

add_h2("3.3  정보 완성도 시스템 (Data Completeness Score)")
add_bullet("금융기관마다 보유 자료 수준이 달라 일부 항목은 미제공 가능 — 이를 투명하게 점수화")
add_bullet("필수(L0) 5점 + 선택(L1) 5개 항목 × 1점 = 총 10점")
add_bullet("매물 카드에 '자료 8/10' 뱃지 표시, 미제공 항목은 ✗로 명시 (빈칸·추정치 금지)")
add_bullet("완성도 9 이상 매물은 프리미엄 노출 무료 → 매각자의 자료 보충 동기 자연 유도")

add_h2("3.4  규제 준수 아키텍처")

add_h3("A. 개인정보 4중 방어선")
add_bullet("[L1] 규칙: 매각자 업로드 시 자동 마스킹(정규식 + NER)")
add_bullet("[L2] 검수: 매각자 1차 확인 + 운영자 2차 이중 검수")
add_bullet("[L3] 격리: 민감 필드는 별도 암호화 테이블(debtor_sensitive_data)에 분리")
add_bullet("[L4] 감사: 모든 L2/L3 접근 로그(pii_access_log) + 이상 패턴 탐지")

add_h3("B. 법적 포지셔닝")
add_bullet("통신판매중개업자로서 중개 서비스만 제공, 채권 매매 당사자 아님")
add_bullet("에스크로는 시중은행 파트너 연동 (전자금융거래법 라이선스 직접 보유 회피)")
add_bullet("신용정보는 신용정보원 비식별 통계 데이터만 활용")
add_bullet("금융위 NPL 매각 규정(공고·경쟁입찰)을 전자화하여 규제당국과 협력 포지션 구축")

add_page_break()

# ════════════════════════════════════════════════════
# 4. 운영 모델 (Operating Model)
# ════════════════════════════════════════════════════
add_h1("Operating Model", number=4)

add_h2("4.1  Revenue Model — 수수료 상한 0.9% 대응")

add_table(
    ["수익원", "요율 / 단가", "연 매출 추정"],
    [
        ["매각자 거래 수수료", "0.5% 기본 ~ 0.9% (옵션)", "300억"],
        ["매수자 거래 수수료", "0.5% 기본 ~ 0.9% (옵션)", "300억"],
        ["구독 (Pro 49만/월)", "전문투자자 500명", "30억"],
        ["구독 (Institutional 490만/월~)", "기관 50곳", "30억"],
        ["에스크로 수수료", "0.3%", "180억"],
        ["실사 대행·계약서", "건당 300~1,500만원", "50억"],
        ["B2B SaaS (마스킹·대시보드)", "연 1~5억 × 20곳", "30억"],
        ["합계", "", "920억"],
    ],
    col_widths=[6, 5, 5]
)

add_h2("4.2  Lock-in Architecture — 플랫폼 방어 장치")

add_h3("공급자 (매각자) 락인")
add_bullet("업계 최저 수수료 0.5% (기존 중개 1~2%, 경매 3~5%)")
add_bullet("AI 자동 평가 + 포트폴리오 대시보드 → 금융기관 내부 툴 대체")
add_bullet("전속 매각 계약: 연 최소 규모 보장 시 수수료 0.3% 할인")
add_bullet("포트폴리오 일괄 상장 + 자동 마스킹 → 운영 비용 절감")

add_h3("수요자 (매수자) 락인")
add_bullet("NDA 비우회 조항: 체결 후 6개월 내 직거래 시 위약금(거래가 5%)")
add_bullet("에스크로 의무화: 자금 흐름이 플랫폼을 통과")
add_bullet("계약 이력·평판·거래 기록이 플랫폼에만 축적")
add_bullet("유사 낙찰 DB·시장 인덱스는 플랫폼 외부에서 재현 불가")

add_h2("4.3  Organization — 필수 조직")
add_table(
    ["Role", "책임", "투입 시기"],
    [
        ["DPO (개인정보보호책임자)", "개인정보 영향평가·사고 대응", "Phase 0"],
        ["준법감시인", "금융 규제 자문", "Phase 0"],
        ["마스킹 검수팀 (2~3명)", "문서 이중 검수", "Phase 3"],
        ["기관영업팀", "저축은행·AMC 전속 계약 체결", "Phase 5"],
        ["CS·분쟁대응", "NDA 위반·이용자 권리 요청 처리", "Phase 4"],
    ],
    col_widths=[5, 7, 4]
)

add_page_break()

# ════════════════════════════════════════════════════
# 5. 실행 로드맵 (Implementation Roadmap)
# ════════════════════════════════════════════════════
add_h1("Implementation Roadmap", number=5)

add_key_takeaway(
    "6 Phase · 30 Business Days — 데이터 계층에서 시작하여 "
    "사용자 인터페이스, 거래 인프라, 관리자 체계 순으로 통합 전개"
)

add_h2("5.1  Phase Overview")
add_table(
    ["Phase", "기간", "핵심 산출물", "의존성"],
    [
        ["Phase 0  데이터 계층", "2일", "타입 확장, 유틸, 데모 데이터, 공통 컴포넌트 11종", "없음 (선행)"],
        ["Phase 1  네비 & 메인", "4일", "5대 메인 네비, 랜딩 페이지 전면 재작성", "Phase 0"],
        ["Phase 2  /exchange 목록·상세", "8일", "티어 게이팅, 완성도 뱃지, 탭 6종, TierGate 적용", "Phase 0·1"],
        ["Phase 3  거래 실행 인프라", "5일", "매각등록, NDA, LOI, 데이터룸, 에스크로", "Phase 2"],
        ["Phase 4  /my 사용자 영역", "4일", "본인인증, KYC, 계약이력, 개인정보, 결제", "Phase 3"],
        ["Phase 5  /admin 관리자", "5일", "마스킹 검수, PII 감사, 기관 관리", "Phase 3"],
        ["Phase 6  통합·상태머신", "2일", "크로스 플로우, deal-state-machine, QA", "전체"],
    ],
    col_widths=[4.5, 1.8, 7, 2.7]
)

add_h2("5.2  Phase 0 — Foundation (2일)")
add_bullet("lib/db-types.ts — DealListing에 outstanding_principal, asking_price, appraisal, registry_summary, rights_structure, lease_summary, provided_fields, data_completeness, access_tier_required, debtor_type 추가")
add_bullet("lib/access-tier.ts — getUserTier, canAccess, maskValue")
add_bullet("lib/data-completeness.ts — calculateCompleteness, getBadgeStyle")
add_bullet("lib/fee-calculator.ts — 상한 0.9% 수수료 계산")
add_bullet("lib/audit-log.ts — logAccess, detectAnomaly")
add_bullet("lib/masking/{engine,pdf,image}.ts — 마스킹 파이프라인 코어")
add_bullet("lib/demo-data/listings.ts — 완성도 분포로 데모 데이터 재생성")
add_bullet("공통 컴포넌트 11종 (TierBadge, TierGate, CompletenessBadge, ProvidedFieldsGrid, MaskedText, FieldRow, KpiCard, DisclosureNotice, AgreementPanel, RegulatoryNotice, AuditBadge)")

add_h2("5.3  Phase 1 — Navigation & Landing (4일)")
add_bullet("5대 메인 네비: 매물 | 거래 | 분석 | 서비스 | 내 정보")
add_bullet("헤더에 사용자 티어 표시 + 업그레이드 CTA")
add_bullet("메인 랜딩 8섹션 재작성: 히어로→가치 3카드→티어 시각화→KPI 미리보기→완성도 설명→기관 로고월→신뢰 배지→CTA")
add_bullet("푸터 규제 고지 강화 (개인정보처리방침·DPO·통신판매중개업 신고번호)")

add_h2("5.4  Phase 2 — Exchange (8일, 핵심)")
add_bullet("목록 페이지: 완성도 슬라이더, 티어 필터, 매물 카드 재디자인, 정렬 4종")
add_bullet("상세 페이지: L0 히어로(KPI 4카드) + 탭 6종(개요·담보권리·자료·실사·Q&A·절차계약)")
add_bullet("각 탭 내부 TierGate로 필드별 L1/L2/L3 게이팅")
add_bullet("사이드 스티키: 현재 티어, 업그레이드 단계, 관심/문의")
add_bullet("L2/L3 표시 시 서버 측 logAccess 호출")

add_h2("5.5  Phase 3 — Transaction Infrastructure (5일)")
add_bullet("/exchange/sell — 4단계 위저드(기본정보→자료 업로드→마스킹 미리보기→제출)")
add_bullet("/exchange/[id]/nda — 전자 NDA 체결 (비밀유지·비우회·위약금 3체크)")
add_bullet("/exchange/[id]/loi — LOI 제출 + 매각자 승인 플로우")
add_bullet("/exchange/[id]/data-room — L3 전용 데이터룸")
add_bullet("/exchange/[id]/escrow — 은행 파트너 계좌 개설 (모의)")

add_h2("5.6  Phase 4 — User Area (4일)")
add_bullet("/my — 티어 대시보드, 업그레이드 단계, 진행중 거래 현황")
add_bullet("/my/verify — 본인인증")
add_bullet("/my/kyc — 전문투자자 인증 (자본시장법 기준)")
add_bullet("/my/agreements — NDA·LOI 이력")
add_bullet("/my/privacy — 개인정보 열람·정정·삭제·동의 철회")
add_bullet("/my/billing — 구독(Basic/Pro/Institutional)·크레딧·수수료 내역")

add_h2("5.7  Phase 5 — Admin (5일)")
add_bullet("/admin/listings 수정 — 심사 상태 컬럼, 완성도 필터, 마스킹 검수 체크")
add_bullet("/admin/masking-queue — 자동 마스킹 결과 이중 검수")
add_bullet("/admin/pii-audit — DPO 전용, 접근 로그·이상 탐지")
add_bullet("/admin/institutions — 기관 등록·전속 계약 관리")
add_bullet("/admin/agreements — NDA/LOI 현황·비우회 위반 플래그")
add_bullet("/admin 대시보드 KPI 교체 — 월 상장, 완결률, 락인율, 침해 0건")

add_h2("5.8  Phase 6 — Integration (2일)")
add_bullet("lib/deal-state-machine.ts — 매물·계약·에스크로 상태 전이 중앙 관리")
add_bullet("매수자 여정·매각자 여정 E2E 연결 검증")
add_bullet("QA 및 완료 기준 체크리스트 점검")

add_page_break()

# ════════════════════════════════════════════════════
# 6. 위험도 및 완화 방안
# ════════════════════════════════════════════════════
add_h1("Risk Matrix & Mitigation", number=6)

add_table(
    ["위험 항목", "가능성", "영향도", "완화 방안"],
    [
        ["자동 마스킹 누락 → 개인정보 노출", "중", "치명", "이중 검수 + 매각자 서약 + 감사 로그"],
        ["금융위·금감원 규제 해석 변동", "중", "대", "준법감시인 상시 자문 + 규제당국 사전 협의"],
        ["매각자 전속계약 확보 지연", "중", "대", "초기 수수료 할인·프리미엄 노출 무상 제공"],
        ["기존 중개업자의 반발", "대", "중", "중개업자와 파트너십(공동 등록) 옵션 제공"],
        ["매수자 허위 인증", "중", "중", "KYC 다단계 + 자본시장법 전문투자자 기준 엄격 적용"],
        ["에스크로 파트너 은행 확보 실패", "소", "대", "복수 은행 병행 협상 + 신탁사 대안"],
        ["비우회 위반 (직거래)", "중", "중", "위약금 5% + 매각자 재상장 차단 + 감사 로그 증거력"],
        ["ISMS-P 인증 지연", "중", "중", "Phase 2부터 준비 착수, 외부 컨설팅 병행"],
    ],
    col_widths=[5, 1.5, 1.5, 8]
)

add_page_break()

# ════════════════════════════════════════════════════
# 7. 기대효과 (Expected Impact)
# ════════════════════════════════════════════════════
add_h1("Expected Impact", number=7)

add_h2("7.1  정량 기대효과 (12개월)")
add_table(
    ["지표", "Baseline", "Target", "증감"],
    [
        ["월 상장 건수", "20건", "100건+", "+400%"],
        ["거래 완결률", "10%", "35%+", "+25%p"],
        ["매각자 락인율", "30%", "85%+", "+55%p"],
        ["평균 자료 완성도", "4.0", "7.5+", "+88%"],
        ["Institutional 고객", "5곳", "50곳+", "+900%"],
        ["연 거래 총액", "500억", "6,000억+", "+1,100%"],
        ["연 매출", "50억", "920억+", "+1,740%"],
        ["개인정보 침해 사고", "미측정", "0건", "완전 통제"],
    ],
    col_widths=[5, 3, 4, 4]
)

add_h2("7.2  정성 기대효과")
add_bullet("규제당국 협력 포지션: 금감원·금융위의 NPL 매각 규정 전자화 파트너로 자리매김")
add_bullet("시장 인프라 지위: NPL 거래의 '코스닥'에 해당하는 공식 인프라로 인식")
add_bullet("데이터 자산화: 비식별 거래 DB가 시세 인덱스·리서치 상품으로 확장 가능")
add_bullet("B2B 확장성: 마스킹 엔진 자체를 금융기관 대상 SaaS로 별도 판매")
add_bullet("해외 진출 기반: 개인정보 보호 체계는 GDPR 호환으로 글로벌 확장 가능")

add_h2("7.3  경쟁 우위 유지 기간")
add_para(
    "본 전략 실행 후 18개월 내 구축되는 데이터·네트워크 효과는 "
    "후발 주자가 3~5년 이내에 복제하기 어려운 진입장벽을 형성한다.",
    size=10, bold=False
)
add_bullet("정보 공개 모델: 3개월 내 복제 가능")
add_bullet("자동 마스킹 기술: 6개월 내 복제 가능")
add_bullet("전속 매각 계약: 12~18개월 확보 필요 (신규 진입자 사실상 불가)")
add_bullet("거래 이력 DB: 24개월 이상 축적 필요 (복제 불가)")
add_bullet("규제당국 신뢰 관계: 수년간 운영 실적 필요 (복제 불가)")

add_page_break()

# ════════════════════════════════════════════════════
# 8. Appendix
# ════════════════════════════════════════════════════
add_h1("Appendix", number=8)

add_h2("A.  핵심 DB 스키마")
add_para("deal_listings (공개 가능 담보·채권 정보)", size=10, bold=True)
add_bullet("outstanding_principal, asking_price, discount_rate — 핵심 재무")
add_bullet("appraisal {value, date, agency, pdf_masked_url}")
add_bullet("registry_summary, rights_structure, lease_summary — L1 요약")
add_bullet("registry_full_url, lease_documents_url — L2 원본")
add_bullet("provided_fields, data_completeness, access_tier_required — 접근 제어")
add_bullet("collateral_region_code (읍/면/동), court_case_masked, debtor_type")

add_para("debtor_sensitive_data (개인정보 분리 암호화)", size=10, bold=True)
add_bullet("encrypted_payload (AES-256), kms_key_id, access_restricted=TRUE")
add_bullet("계약 체결 후 법적 채권양도 절차에서만 복호화")

add_para("pii_access_log (감사 로그)", size=10, bold=True)
add_bullet("user_id, target_table, target_id, access_type, purpose, ip_address, accessed_at")

add_para("deal_agreements (NDA/LOI 이력)", size=10, bold=True)
add_bullet("listing_id, user_id, agreement_type, signed_at, signature_hash, non_circumvention_until")

add_h2("B.  필수 인허가·신고")
add_table(
    ["항목", "근거 법령", "대응 방안"],
    [
        ["통신판매중개업 신고", "전자상거래법", "Phase 0 즉시 신고"],
        ["개인정보처리자 지정", "개인정보보호법", "DPO 선임·개인정보 영향평가"],
        ["ISMS-P 인증", "정보통신망법", "Phase 2 착수·Phase 4 완료 목표"],
        ["에스크로", "전자금융거래법", "시중은행 파트너 연동 (직접 보유 회피)"],
        ["전문투자자 확인", "자본시장법", "KYC 단계 적용"],
        ["신용정보 이용", "신용정보법", "비식별 통계만 활용"],
    ],
    col_widths=[4, 4, 8]
)

add_h2("C.  완료 기준 체크리스트")

add_h3("기능 요건")
add_bullet("L0~L3 티어 게이팅이 모든 민감 필드에 적용")
add_bullet("매물 카드·상세에 완성도 뱃지 표시")
add_bullet("매각자 업로드 → 자동 마스킹 → 검수 → 공개 플로우 동작")
add_bullet("NDA·LOI 체결 후 티어 자동 해금")
add_bullet("L2/L3 접근이 pii_access_log에 기록")
add_bullet("수수료 상한 0.9% 자동 적용")
add_bullet("관리자 마스킹 검수·PII 감사 대시보드 동작")

add_h3("디자인 요건")
add_bullet("모든 다크 섹션 인라인 C.bgN 사용")
add_bullet("dark: variant·shadcn·card-elevated·금지 import 전무")
add_bullet("티어·완성도·규제 색상 일관 적용")
add_bullet("공통 컴포넌트 11종이 전 페이지에서 재사용")

add_h3("규제 요건")
add_bullet("개인정보처리방침·약관·NDA 템플릿 페이지 존재")
add_bullet("이용자 권리(열람·정정·삭제) 관리 페이지 동작")
add_bullet("DPO 연락처 푸터 노출")
add_bullet("규제 고지 문구 상세·데이터룸·NDA 페이지에 일괄 표시")

# 문서 끝
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— END OF REPORT —")
set_run(run, size=9, color=GRAY_MID, font=FONT_EN)

# 저장
import os
out_dir = r"C:\Users\82106\Desktop\nplatform\docs"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "NPLatform_Strategic_Transformation_Report_v4.docx")
doc.save(out_path)
print(f"SAVED: {out_path}")
print(f"SIZE: {os.path.getsize(out_path):,} bytes")
