# -*- coding: utf-8 -*-
"""Combine NPLatform strategy markdown docs into a single Word file."""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = r"C:\Users\82106\Desktop\nplatform\docs"
OUT = r"C:\Users\82106\Desktop\nplatform\docs\NPLatform_종합기획서.docx"

FILES = [
    ("NPLatform_Strategy_McKinsey.md",                "Part A — McKinsey 스타일 전략 기획서"),
    ("NPLatform_v6_Strategy_and_Engineering.md",      "Part B — 통합 기획서 (Strategy × Engineering v6)"),
    ("NPLatform_Connected_Experience_Plan.md",        "Part C — Connected Experience 설계서"),
    ("NPLatform_Global_Benchmark_and_Roadmap.md",     "Part D — 글로벌 벤치마크 & 12개월 로드맵"),
]

# ─── Style helpers ───────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ─── Inline parsing (bold/italic/code) ───────────────────────────────

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')

def add_inline_runs(paragraph, text, base_size=10.5, base_color=None):
    """Parse inline markdown and add runs."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.text = part
        if run.font.size is None:
            run.font.size = Pt(base_size)
        if base_color and run.font.color.rgb is None:
            run.font.color.rgb = base_color


# ─── Block parsing ───────────────────────────────────────────────────

def parse_table(lines, i):
    """Parse a markdown table starting at line i. Returns (rows, new_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [''] * (n_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text, base_size=9)
            if r_idx == 0:
                set_cell_bg(cell, '1F4E78')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)


def render_md(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code:
                # close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # gray bg via paragraph shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F4F4F4')
                pPr.append(shd)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            elif level == 2:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                # bottom border
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '2E74B5')
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif level == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(12.5)
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5C)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-=*_]{3,}$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table
        if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            doc.add_paragraph()  # spacer
            continue

        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_inline_runs(p, ' '.join(quote_lines), base_size=10.5,
                            base_color=RGBColor(0x4A, 0x4A, 0x4A))
            for run in p.runs:
                run.italic = True
            # left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '2E74B5')
            pBdr.append(left)
            pPr.append(pBdr)
            continue

        # Bullet / numbered list
        m = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
        if m:
            indent_spaces = len(m.group(1))
            marker = m.group(2)
            text = m.group(3)
            text = re.sub(r'^\[[ xX]\]\s*', '', text)
            p = doc.add_paragraph(style='List Bullet' if marker in '-*+' else 'List Number')
            p.paragraph_format.left_indent = Cm(0.6 + indent_spaces * 0.3)
            p.paragraph_format.space_after = Pt(2)
            # clear default text
            for r in list(p.runs):
                r.text = ''
            add_inline_runs(p, text, base_size=10.5)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, line, base_size=10.5)
        i += 1


def main():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    rpr = style.element.rPr
    rFonts = rpr.rFonts if rpr is not None else None
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('NPLatform')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('한국 부실채권·부동산 통합 거래 플랫폼')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('종합 기획서')
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run('Strategy · Product · Engineering · Design · Operations')
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(200)
    run = p.add_run('2026-04-07 · v1.0\nStrictly Confidential — Draft for Discussion')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_page_break(doc)

    # TOC of parts
    p = doc.add_paragraph()
    run = p.add_run('Contents')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    for idx, (_, title) in enumerate(FILES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    add_page_break(doc)

    # Render each part
    for idx, (filename, title) in enumerate(FILES):
        path = os.path.join(DOCS_DIR, filename)
        if not os.path.exists(path):
            print(f'[skip] {filename} not found')
            continue
        with open(path, 'r', encoding='utf-8') as f:
            md = f.read()

        # Part divider
        if idx > 0:
            add_page_break(doc)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(filename)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        add_page_break(doc)

        render_md(doc, md)

    doc.save(OUT)
    print(f'[done] saved to {OUT}')
    print(f'[size] {os.path.getsize(OUT) / 1024:.1f} KB')


if __name__ == '__main__':
    main()
